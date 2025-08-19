import re
import pdfplumber
from docx import Document
from fastapi import UploadFile

def formatting_score_and_suggestions(resume: UploadFile) -> tuple:
    """
    Evaluates resume formatting for ATS-friendliness and offers suggestions.
    Args:
        resume (UploadFile): The uploaded resume file.
    Returns:
        tuple: (score, list of feedback strings)
    """
    feedback = []
    max_points = 10
    deductions = 0
    STANDARD_FONTS = {'arial', 'calibri', 'times new roman'}
    DATE_PATTERNS = [
        r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}",
        r"\d{1,2}/\d{4}",
        r"\d{4}",
        r"\d{1,2}-\d{4}",
        r"(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}"
    ]
    # More permissive filename patterns
    PROBLEMATIC_CHARS = re.compile(r"[<>:\"/\\|?*]")  # Characters that cause file system issues
    GENERIC_NAMES = {'resume.pdf', 'resume.docx', 'cv.pdf', 'cv.docx', 'untitled.pdf', 'untitled.docx'}
    
    resume.file.seek(0)
    resume_name = resume.filename.lower() if resume.filename else ""

    # ----------- FILE NAME CHECK (Minor deduction) -----------
    filename_issues = []
    
    # Check for problematic characters
    if PROBLEMATIC_CHARS.search(resume.filename or ""):
        filename_issues.append("contains special characters that may cause issues")
    
    # Check for generic naming
    if resume_name in GENERIC_NAMES:
        filename_issues.append("use a more specific name like 'YourName_Resume.pdf'")
    
    # Check for overly long names (some systems have limits)
    if len(resume.filename or "") > 100:
        filename_issues.append("filename is too long (over 100 characters)")
    
    # Check for meaningful content
    name_without_ext = resume_name.rsplit('.', 1)[0] if '.' in resume_name else resume_name
    if len(name_without_ext.replace('_', '').replace('-', '').replace(' ', '')) < 3:
        filename_issues.append("filename should be more descriptive")
    
    if filename_issues:
        deductions += 3  # Even smaller deduction
        issue_text = ", ".join(filename_issues)
        feedback.append(f"• Consider renaming your file - it {issue_text}. Example: 'YourName_Resume.pdf'")

    # ----------- PDF HANDLING -----------
    if resume.content_type == 'application/pdf' or resume_name.lower().endswith('.pdf'):
        with pdfplumber.open(resume.file) as pdf:
            fonts_found = set()
            long_paragraphs = 0
            date_format_issues = 0
            has_tables = False
            has_images = False

            for page in pdf.pages:
                # Check for images
                if page.images:
                    has_images = True
                
                # Check for tables
                if page.extract_tables():
                    has_tables = True

                for char in page.chars:
                    font_name = char.get('fontname', '').lower()
                    if 'arial' in font_name:
                        fonts_found.add('arial')
                    elif 'calibri' in font_name:
                        fonts_found.add('calibri')
                    elif 'times' in font_name:
                        fonts_found.add('times new roman')
                    else:
                        fonts_found.add(font_name.split('+')[-1])

                text = page.extract_text() or ""
                paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                for para in paragraphs:
                    # More lenient paragraph length check
                    if len(para.split()) > 50:  # Increased threshold
                        long_paragraphs += 1
                    
                    # Better date pattern matching
                    dates = re.findall(r'\b\w*(\d{4}|\d{1,2}/\d{4}|[A-Za-z]{3,9}\s+\d{4}|\d{1,2}-\d{4})\w*\b', para)
                    for date in dates:
                        if not any(re.search(pat, str(date)) for pat in DATE_PATTERNS):
                            date_format_issues += 1

            # Font check - more lenient
            non_standard_fonts = fonts_found - STANDARD_FONTS
            if non_standard_fonts and len(non_standard_fonts) > 2:  # Only penalize if many non-standard fonts
                deductions += 2
                feedback.append("Use standard fonts like Arial, Calibri, or Times New Roman for better ATS compatibility.")

            # Table check - major deduction
            if has_tables:
                deductions += 3
                feedback.append("Avoid tables – ATS tools often can't read text inside tables properly.")

            # Image check - major deduction
            if has_images:
                deductions += 3
                feedback.append("Avoid using images – ATS systems don't parse images and may skip important info.")

            # Long paragraphs - minor deduction
            if long_paragraphs > 3:  # Only penalize if many long paragraphs
                deductions += 1
                feedback.append("Break long paragraphs into bulleted points for improved readability.")

            # Date format issues - minor deduction
            if date_format_issues > 5:  # Only penalize if many inconsistent dates
                deductions += 1
                feedback.append("Unify your date formats consistently (e.g., Jan 2020 – Mar 2022).")

    # ----------- DOCX HANDLING -----------
    elif resume.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or resume_name.lower().endswith('.docx'):
        doc = Document(resume.file)
        resume.file.seek(0)
        fonts_found = set()
        long_paragraphs = 0
        date_format_issues = 0
        header_footer_text_found = False

        def get_effective_font(run, paragraph, doc):
            if run.font.name:
                return run.font.name
            if run.style and hasattr(run.style, "font") and run.style.font.name:
                return run.style.font.name
            if paragraph.style and hasattr(paragraph.style, "font") and paragraph.style.font.name:
                return paragraph.style.font.name
            try:
                normal_style = doc.styles['Normal']
                if hasattr(normal_style, "font") and normal_style.font.name:
                    return normal_style.font.name
            except Exception:
                pass
            return None

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                font = get_effective_font(run, paragraph, doc)
                if font:
                    fonts_found.add(font.strip().lower())
            
            # More lenient paragraph length check
            if len(paragraph.text.split()) > 50:  # Increased threshold
                long_paragraphs += 1
            
            # Better date pattern matching
            dates = re.findall(r'\b\w*(\d{4}|\d{1,2}/\d{4}|[A-Za-z]{3,9}\s+\d{4}|\d{1,2}-\d{4})\w*\b', paragraph.text)
            for date in dates:
                if not any(re.search(pat, str(date)) for pat in DATE_PATTERNS):
                    date_format_issues += 1

        # Font check - more lenient
        non_standard_fonts = fonts_found - STANDARD_FONTS
        if non_standard_fonts and len(non_standard_fonts) > 2:  # Only penalize if many non-standard fonts
            deductions += 2
            feedback.append("Stick to safe fonts like Arial, Calibri, or Times New Roman — these are ATS-compatible.")

        # Table check - major deduction
        if doc.tables:
            deductions += 3
            feedback.append("Avoid using tables — they may not be parsed correctly by ATS systems.")

        # Image check - major deduction  
        if hasattr(doc, 'inline_shapes') and doc.inline_shapes:
            deductions += 3
            feedback.append("Remove images — they are skipped by resume parsers.")

        # Long paragraphs - minor deduction
        if long_paragraphs > 3:  # Only penalize if many long paragraphs
            deductions += 1
            feedback.append("Consider converting long paragraphs into concise bullet points.")

        # Date format issues - minor deduction
        if date_format_issues > 5:  # Only penalize if many inconsistent dates
            deductions += 1
            feedback.append("Use consistent formatting for dates, like 'Jan 2020 – Mar 2022'.")

        # Header/footer check - minor deduction
        try:
            for section in doc.sections:
                header_text = " ".join([p.text for p in section.header.paragraphs]).strip()
                footer_text = " ".join([p.text for p in section.footer.paragraphs]).strip()
                if header_text or footer_text:
                    header_footer_text_found = True
                    break
        except Exception:
            pass

        if header_footer_text_found:
            deductions += 0.5  # Minor deduction
            feedback.append("Important content should be in the body — ATS may skip headers/footers.")

    # ----------- FINAL SCORING -----------
    formatting_score = max(max_points - deductions, 0)
    
    # Ensure integer score for good presentation
    formatting_score = int(round(formatting_score))
    
    # Add contextual feedback based on score ranges
    if not feedback:
        feedback.append("✅ Excellent! Your resume formatting is clean and ATS-friendly.")
    elif formatting_score >= 8:
        feedback.insert(0, "✅ Great formatting! Just minor tweaks suggested.")
    elif formatting_score >= 6:
        feedback.insert(0, "✅ Good formatting overall with some areas for improvement.")
    elif formatting_score >= 4:
        feedback.insert(0, "⚠️ Decent formatting but several issues may affect ATS parsing.")
    else:
        feedback.insert(0, "❌ Multiple formatting issues detected that will likely affect ATS parsing.")
    
    return formatting_score, feedback